from io import BytesIO
from typing import Dict, Union
from docx import Document


def _parse_document(doc: Document) -> Dict[str, str]:
    """
    Extrai seções 'relatorio', 'fundamentacao' e 'dispositivo' de um Document do python-docx.
    Retorna um dicionário com as chaves:
      - 'relatorio'
      - 'fundamentacao'
      - 'dispositivo'
    Cada valor é o texto concatenado da respectiva seção.
    """
    secoes = {"relatorio": "", "fundamentacao": "", "dispositivo": ""}
    secao_atual: Union[str, None] = None

    for para in doc.paragraphs:
        texto = para.text.strip()
        if not texto:
            continue

        lower = texto.lower()
        # Identifica cabeçalhos de seção
        if lower.startswith("relatório") or lower.startswith("relatorio"):
            secao_atual = "relatorio"
            continue
        if lower.startswith("fundamentação") or lower.startswith("fundamentacao"):
            secao_atual = "fundamentacao"
            continue
        if lower.startswith("dispositivo"):
            secao_atual = "dispositivo"
            continue

        # Adiciona texto à seção atual
        if secao_atual:
            secoes[secao_atual] += texto + "\n"

    return secoes


def parse_docx_file(path: str) -> Dict[str, str]:
    """
    Carrega um arquivo .docx do sistema de arquivos e extrai as seções.

    :param path: caminho para o arquivo .docx
    :return: dict com as seções extraídas
    """
    doc = Document(path)
    return _parse_document(doc)


def parse_docx_bytes(data: bytes) -> Dict[str, str]:
    """
    Recebe bytes de um arquivo .docx (por exemplo, UploadFile.read()) e extrai as seções.

    :param data: conteúdo binário do .docx
    :return: dict com as seções extraídas
    """
    doc = Document(BytesIO(data))
    return _parse_document(doc)
