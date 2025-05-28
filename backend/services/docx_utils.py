import os
import zipfile
from datetime import datetime
from typing import List, Dict, Optional
from docx import Document
from docx.shared import Inches


def salvar_sentenca_como_docx(
    relatorio: str,
    fundamentacao_dispositivo: str,
    arquivo_path: str,
    numero_processo: Optional[str] = None
):
    """
    Salva uma sentença completa em formato DOCX com formatação adequada
    
    Args:
        relatorio: O relatório do processo
        fundamentacao_dispositivo: A fundamentação e dispositivo gerados
        arquivo_path: Caminho onde salvar o arquivo
        numero_processo: Número do processo (opcional)
    """
    doc = Document()
    
    # Configurações de página
    section = doc.sections[0]
    section.page_height = Inches(11.69)  # A4
    section.page_width = Inches(8.27)    # A4
    section.left_margin = Inches(1.18)   # 3cm
    section.right_margin = Inches(0.79)  # 2cm
    section.top_margin = Inches(0.79)    # 2cm
    section.bottom_margin = Inches(0.79) # 2cm
    
    # Cabeçalho do documento
    if numero_processo:
        heading = doc.add_heading(f"SENTENÇA", level=1)
        heading.alignment = 1  # Centralizado
        
        # Número do processo
        processo_para = doc.add_paragraph()
        processo_para.alignment = 1  # Centralizado
        run = processo_para.add_run(f"Processo nº {numero_processo}")
        run.bold = True
    else:
        heading = doc.add_heading("SENTENÇA", level=1)
        heading.alignment = 1  # Centralizado
    
    # Quebra de linha
    doc.add_paragraph()
    
    # Adiciona o relatório
    if relatorio and relatorio.strip():
        doc.add_heading("RELATÓRIO", level=2)
        
        # Processa o relatório preservando formatação
        paragrafos_relatorio = relatorio.split('\n\n')
        for paragrafo in paragrafos_relatorio:
            if paragrafo.strip():
                doc.add_paragraph(paragrafo.strip())
        
        doc.add_paragraph()  # Espaço
    
    # Adiciona a fundamentação e dispositivo
    if fundamentacao_dispositivo and fundamentacao_dispositivo.strip():
        # Se não há seções explícitas, assume que é tudo fundamentação seguida de dispositivo
        if "FUNDAMENTAÇÃO" not in fundamentacao_dispositivo.upper() and "DISPOSITIVO" not in fundamentacao_dispositivo.upper():
            doc.add_heading("FUNDAMENTAÇÃO E DISPOSITIVO", level=2)
        
        # Processa o texto preservando formatação
        paragrafos = fundamentacao_dispositivo.split('\n\n')
        for paragrafo in paragrafos:
            if paragrafo.strip():
                # Verifica se é um cabeçalho de seção
                texto_limpo = paragrafo.strip()
                if (texto_limpo.upper().startswith('FUNDAMENTAÇÃO') or 
                    texto_limpo.upper().startswith('DISPOSITIVO') or
                    texto_limpo.upper().startswith('MÉRITO')):
                    doc.add_heading(texto_limpo, level=2)
                else:
                    doc.add_paragraph(texto_limpo)
    
    # Salva o documento
    doc.save(arquivo_path)


def salvar_docs_referencia(docs: List[Dict], arquivo_zip_path: str):
    """
    Salva os documentos de referência em um arquivo ZIP
    
    Args:
        docs: Lista de documentos com as seções
        arquivo_zip_path: Caminho onde salvar o ZIP
    """
    try:
        with zipfile.ZipFile(arquivo_zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for i, doc in enumerate(docs, 1):
                # Cria um documento DOCX para cada referência
                doc_temp = Document()
                
                # Título do documento
                doc_id = doc.get('id', f'documento_{i}')
                doc_temp.add_heading(f"Documento de Referência {i}", level=1)
                doc_temp.add_paragraph(f"ID: {doc_id}")
                doc_temp.add_paragraph()
                
                # Adiciona as seções se existirem
                if doc.get('relatorio'):
                    doc_temp.add_heading("RELATÓRIO", level=2)
                    paragrafos = doc['relatorio'].split('\n\n')
                    for p in paragrafos:
                        if p.strip():
                            doc_temp.add_paragraph(p.strip())
                    doc_temp.add_paragraph()
                
                if doc.get('fundamentacao'):
                    doc_temp.add_heading("FUNDAMENTAÇÃO", level=2)
                    paragrafos = doc['fundamentacao'].split('\n\n')
                    for p in paragrafos:
                        if p.strip():
                            doc_temp.add_paragraph(p.strip())
                    doc_temp.add_paragraph()
                
                if doc.get('dispositivo'):
                    doc_temp.add_heading("DISPOSITIVO", level=2)
                    paragrafos = doc['dispositivo'].split('\n\n')
                    for p in paragrafos:
                        if p.strip():
                            doc_temp.add_paragraph(p.strip())
                
                # Adiciona informações de score se disponíveis
                if doc.get('score') is not None or doc.get('rerank_score') is not None:
                    doc_temp.add_paragraph()
                    doc_temp.add_heading("INFORMAÇÕES DE RELEVÂNCIA", level=3)
                    if doc.get('score') is not None:
                        doc_temp.add_paragraph(f"Score de similaridade: {doc['score']:.4f}")
                    if doc.get('rerank_score') is not None:
                        doc_temp.add_paragraph(f"Score de re-ranking: {doc['rerank_score']:.4f}")
                
                # Salva temporariamente e adiciona ao ZIP
                # Remove caracteres problemáticos do doc_id
                doc_id_limpo = doc_id.replace('/', '_').replace('\\', '_')
                temp_filename = f"referencia_{i:02d}_{doc_id_limpo}.docx"
                temp_path = f"/tmp/{temp_filename}"
                
                try:
                    doc_temp.save(temp_path)
                    zipf.write(temp_path, temp_filename)
                finally:
                    # Remove arquivo temporário
                    if os.path.exists(temp_path):
                        os.remove(temp_path)
                        
    except Exception as e:
        print(f"Erro ao criar ZIP de referências: {e}")
        # Cria um ZIP vazio em caso de erro para não quebrar o fluxo
        with zipfile.ZipFile(arquivo_zip_path, 'w') as zipf:
            zipf.writestr("erro.txt", f"Erro ao processar documentos de referência: {str(e)}")


def criar_docx_simples(conteudo: str, titulo: str = "Documento") -> bytes:
    """
    Cria um documento DOCX simples em memória e retorna os bytes
    
    Args:
        conteudo: Texto do documento
        titulo: Título do documento
        
    Returns:
        bytes: Conteúdo do arquivo DOCX
    """
    from io import BytesIO
    
    doc = Document()
    doc.add_heading(titulo, level=1)
    
    # Processa o conteúdo
    paragrafos = conteudo.split('\n\n')
    for paragrafo in paragrafos:
        if paragrafo.strip():
            doc.add_paragraph(paragrafo.strip())
    
    # Salva em memória
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()