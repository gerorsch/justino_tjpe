import streamlit as st
import requests
import os
import time
from io import BytesIO
from docx import Document
import re
from datetime import datetime
import json
from sseclient import SSEClient
from dotenv import load_dotenv

# --- Configuração Inicial ---
st.set_page_config(page_title="Justino — Assessor Digital TJPE", page_icon="⚖️", layout="wide")
load_dotenv()
API_URL = os.getenv("API_URL", "http://localhost:8001") # URL do seu backend FastAPI

# --- Funções Auxiliares da Aplicação (do seu código original) ---

def limpar_relatorio(texto_bruto):
    """
    Remove tags, formatação e outros elementos indesejados do relatório
    """
    if not texto_bruto:
        return ""
    original_len = len(texto_bruto)
    texto = re.sub(r'\[TextBlock\([^]]*\)\]', '', texto_bruto)
    texto = re.sub(r'TextBlock\([^)]*\)', '', texto)
    if texto.startswith('data:'):
        texto = texto[5:].strip()
    texto = re.sub(r'citations=None,\s*text=', '', texto)
    texto = re.sub(r'citations=[^,]*,\s*text=', '', texto)
    texto = re.sub(r"type='text'", '', texto)
    texto = texto.strip('"\'')
    texto = texto.replace("'", "'")
    texto = texto.replace('\\"', '"')
    texto = texto.replace('\\n', '\n')
    texto = texto.replace('\\t', '\t')
    texto = texto.replace('\\r', '')
    texto = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', texto)
    linhas = texto.split('\n')
    linhas_limpas = [linha.strip() for linha in linhas if linha.strip() and not re.match(r'^[\[\](),.;:\'"\\/-]+$', linha.strip())]
    relatorio_final = '\n\n'.join(linhas_limpas)
    relatorio_final = re.sub(r'\n{3,}', '\n\n', relatorio_final)
    resultado = relatorio_final.strip()
    if len(resultado) < (original_len * 0.1) and original_len > 100:
        return texto_bruto.strip()
    return resultado

def extrair_numero_processo(texto):
    if not texto:
        return None
    padroes = [
        r'\b\d{7}-\d{2}\.\d{4}\.\d{1}\.\d{2}\.\d{4}\b',
        r'\b\d{10}-\d{2}\.\d{4}\.\d{1}\.\d{2}\.\d{4}\b',
        r'\b\d{4}\.\d{2}\.\d{6}-\d{1}\b',
        r'(?:número|processo|autos)(?:\s*:?\s*|\s+n[º°]?\.?\s*)(\d{7}-\d{2}\.\d{4}\.\d{1}\.\d{2}\.\d{4})',
        r'(?:processo|autos)(?:\s+n[º°]?\.?\s*|\s+)(\d+[-\.\d]+)',
    ]
    for i, padrao in enumerate(padroes):
        matches = re.findall(padrao, texto, re.IGNORECASE)
        if matches:
            numero = matches[0] if isinstance(matches[0], str) else matches[0]
            if re.match(r'\d{7}-\d{2}\.\d{4}\.\d{1}\.\d{2}\.\d{4}', numero) or re.match(r'\d{10}-\d{2}\.\d{4}\.\d{1}\.\d{2}\.\d{4}', numero) or len(numero) > 10:
                return numero
    return None

def gerar_nome_arquivo_sentenca(numero_processo=None):
    """Gera um nome de arquivo inteligente para a sentença."""
    if numero_processo:
        numero_limpo = re.sub(r'[\D]', '', numero_processo)
        return f"sentenca_{numero_limpo}_{datetime.now().strftime('%Y%m%d')}.docx"
    return f"sentenca_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

# --- Lógica de Autenticação e Sessão ---

def init_session_state():
    """Inicializa todas as variáveis de estado da sessão."""
    defaults = {
        "authenticated": False, "jwt_token": None, "user_email": None, "login_step": "email",
        "relatorio": None, "relatorio_processado": False, "sentenca_texto": None,
        "sentenca_processada": False, "sentenca_bytes": None, "referencias_bytes": None,
        "numero_processo": None, "confirming_clear": False, "processed_file_id": None, "generating_sentence": False
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value

def clear_working_data():
    """Reseta os dados de trabalho (relatório, sentença, etc.) na sessão."""
    st.session_state.relatorio = None
    st.session_state.relatorio_processado = False
    st.session_state.sentenca_texto = None
    st.session_state.sentenca_processada = False
    st.session_state.sentenca_bytes = None
    st.session_state.referencias_bytes = None
    st.session_state.numero_processo = None
    st.session_state.confirming_clear = False
    st.session_state.processed_file_id = None
    st.session_state.generating_sentence = False
    
    st.toast("Dados limpos. Pronto para um novo processo!", icon="✨")
    time.sleep(1)
    st.rerun()

def show_login_page():
    """Exibe a interface de login."""
    st.markdown("""
    <style>
    .login-box {
        background-color: #273449; border-radius: 10px; padding: 2rem;
        box-shadow: 0 4px 12px rgba(0,0,0,0.3);
    }
    .login-header { text-align: center; margin-bottom: 1.5rem; }
    .login-header h1 { font-size: 2.2rem; color: #f1f5f9; margin: 0; }
    .login-header p { color: #cbd5e1; margin-top: 0.5rem; }
    </style>
    """, unsafe_allow_html=True)

    _, login_col, _ = st.columns([1, 1.5, 1])

    with login_col:
        st.markdown("<div class='login-box'>", unsafe_allow_html=True)

        if st.session_state.login_step == "email":
            st.markdown("<div class='login-header'><h1>⚖️ Justino Digital</h1><p>Acesse com seu e-mail institucional.</p></div>", unsafe_allow_html=True)
            with st.form("email_form"):
                email_input = st.text_input("E-mail", placeholder="seu.nome@tjpe.jus.br", label_visibility="collapsed")
                if st.form_submit_button("Enviar código de acesso", use_container_width=True):
                    if not email_input:
                        st.error("Por favor, insira um e-mail.")
                    else:
                        try:
                            with st.spinner("Solicitando código..."):
                                response = requests.post(f"{API_URL}/auth/request-code", json={"email": email_input})
                            if response.status_code == 200:
                                st.session_state.user_email = email_input
                                st.session_state.login_step = "code"
                                st.success(f"Código enviado para {email_input}.")
                                time.sleep(2); st.rerun()
                            else:
                                error_data = response.json()
                                detail = error_data.get('detail', 'Erro desconhecido.')
                                if response.status_code == 401:
                                    st.error(f"Acesso negado: {detail} Contate george.queiroz@tjpe.jus.br.")
                                else:
                                    st.error(f"Erro: {detail}")
                        except requests.ConnectionError:
                            st.error("Erro de conexão com o servidor.")
                        except Exception as e:
                            st.error(f"Ocorreu um erro: {e}")

        elif st.session_state.login_step == "code":
            st.markdown(f"<div class='login-header'><h1>🔐 Verifique seu Acesso</h1><p>Digite o código enviado para <strong>{st.session_state.user_email}</strong>.</p></div>", unsafe_allow_html=True)
            with st.form("code_form"):
                code_input = st.text_input("Código de Acesso", max_chars=6, placeholder="******", label_visibility="collapsed")
                if st.form_submit_button("Validar Código", use_container_width=True):
                    if not code_input:
                        st.error("Por favor, insira o código.")
                    else:
                        try:
                            with st.spinner("Validando..."):
                                response = requests.post(f"{API_URL}/auth/verify-code", json={"email": st.session_state.user_email, "code": code_input})
                            if response.status_code == 200:
                                st.session_state.jwt_token = response.json().get("access_token")
                                st.session_state.authenticated = True
                                st.success("Login realizado com sucesso!")
                                time.sleep(1); st.rerun()
                            else:
                                st.error(f"Erro: {response.json().get('detail', 'Código inválido.')}")
                        except requests.ConnectionError:
                            st.error("Erro de conexão com o servidor.")
                        except Exception as e:
                            st.error(f"Ocorreu um erro: {e}")
        st.markdown("</div>", unsafe_allow_html=True)

# --- Aplicação Principal ---

def main_app():
    with st.sidebar:
        st.title("📋 Instruções")
        st.info(f"Sessão iniciada como: **{st.session_state.user_email}**")
        if st.button("Sair"):
            for key in list(st.session_state.keys()): del st.session_state[key]
            st.rerun()
        st.markdown("---")
        st.markdown("### 🚀 Como usar o Justino")
            
        st.markdown("#### **1. Geração do Relatório**")
        st.markdown("""
        - Baixe o processo do PJe em ordem CRESCENTE
        - Faça o upload do processo em PDF (máx. 200MB)
        - Clique em **"Gerar Relatório"**
        - Aguarde o processamento completo
        - Baixe o relatório em formato DOC ou DOCX
        """)
            
        st.markdown("#### **2. Geração da Sentença**")
        st.markdown("""
        - **Instruções Adicionais** (opcional): 
          - Orientações específicas para a sentença
          - Pontos que devem ser destacados
          - Particularidades do caso
            
        - **Documentos de Referência** (opcional):
          - Adicione seus modelos/sentenças similares no formato DOCX
        """)
            
        st.markdown("#### **📁 Formatos Suportados**")
        st.markdown("""
        - **Upload**: PDF (processos)
        - **Referências**: DOCX (sentenças)
        - **Download**: DOCX (relatórios e sentenças)
        """)
            
        st.markdown("#### **⚠️ Dicas Importantes**")
        st.info("""
        🔸 **Qualidade do PDF**: Certifique-se de que o texto do PDF seja legível. Processos migrados não funcionam muito bem.
            
        🔸 **Documentos de Referência**: Inclua modelos/sentenças similares caso exista.
            
        🔸 **Instruções Específicas**: Seja claro sobre aspectos particulares do caso.
            
        🔸 **Revisão Manual**: Sempre revise a sentença gerada antes do uso.
        """)         
            
        st.markdown("#### **📞 Suporte**")
        st.markdown("""
        Para dúvidas, sugestões ou críticas:
        - **Email**: george.queiroz@tjpe.jus.br
        - **Versão**: BETA v2.2 (Cadastro de usuários)
        """)
            
        # Status do sistema
        st.markdown("---")
        st.markdown("#### **📊 Status do Sistema**")
        try:
            resp = requests.get(f"{API_URL}/health", timeout=None)
            if resp.status_code == 200:
                st.success("🟢 Sistema Online")
            else:
                st.warning("🟡 Sistema com Problemas")
        except:
            st.error("🔴 Sistema Offline")
    
    _, main_col, _ = st.columns([1, 3, 1])

    with main_col:
        st.title("⚖️ Justino — Assessor Digital TJPE")
        st.markdown("---")
        st.header("1. Extração do Relatório")
        
        uploaded_pdf = st.file_uploader("📎 Envie um processo em PDF", type=["pdf"], key="uploader_pdf")

        if uploaded_pdf and st.session_state.relatorio_processado and uploaded_pdf.file_id != st.session_state.processed_file_id:
            st.warning("**Atenção:** Já existe um relatório processado. Fazer um novo upload irá apagar os dados atuais.")
            col1, col2, col3 = st.columns(3)
            with col3:
                if st.button("Novo upload", use_container_width=True, type="primary"):
                    clear_working_data()

        if uploaded_pdf and not st.session_state.relatorio_processado:
            col1, col2, col3 = st.columns(3)
            with col3:
                if st.button("Gerar Relatório", key="btn_extrair", use_container_width=True):
                    status_placeholder = col1.empty()
                    with status_placeholder:
                        with st.spinner("Extraindo relatório... O processo pode demorar alguns minutos."):
                            files = {"pdf": (uploaded_pdf.name, uploaded_pdf.getvalue(), "application/pdf")}
                            try:
                                resp_direct = requests.post(f"{API_URL}/processar", files=files, timeout=600)
                                if resp_direct.status_code == 200:
                                    result = resp_direct.json()
                                    if 'relatorio' in result and result['relatorio']:
                                        relatorio_limpo = limpar_relatorio(result['relatorio'])
                                        st.session_state.relatorio = relatorio_limpo
                                        st.session_state.numero_processo = extrair_numero_processo(relatorio_limpo)
                                        st.session_state.relatorio_processado = True
                                        st.session_state.processed_file_id = uploaded_pdf.file_id
                                        st.rerun()
                                    else:
                                        st.error("❌ API retornou uma resposta vazia.")
                                else:
                                    st.error(f"❌ Erro {resp_direct.status_code}: {resp_direct.text}")
                            except Exception as e:
                                st.error(f"❌ Erro inesperado durante a extração: {e}")

        if st.session_state.relatorio_processado:
            st.success(f"📄 Relatório extraído! Processo: **{st.session_state.numero_processo or 'Não identificado'}**")
            with st.expander("📄 Visualizar Relatório Extraído", expanded=True):
                st.text_area("Conteúdo:", value=st.session_state.relatorio, height=300, disabled=True)
            buffer = BytesIO()
            doc = Document()
            doc.add_heading("Relatório Extraído", level=1)
            doc.add_paragraph(st.session_state.relatorio)
            doc.save(buffer)
            buffer.seek(0)
            nome_arquivo_relatorio = f"relatorio_{st.session_state.numero_processo or 'sem_numero'}.docx"
            col1, col2, col3 = st.columns(3)
            with col3:
                st.download_button(label="Baixar Relatório", data=buffer.getvalue(), file_name=nome_arquivo_relatorio, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
            st.markdown("---")
        
        st.header("2. Geração da Sentença")
        if not st.session_state.relatorio_processado:
            st.warning("⚠️ Extraia um relatório para continuar.")
        else:
            instrucoes = st.text_area("Instruções Adicionais", placeholder=("Ex: quantificar danos morais, valorar determinada prova, determinar a procedência ou improcedência do pedido, etc."))
            
            arquivos_ref = st.file_uploader("Documentos de Referência (Opcional)", type=["docx", "doc"], accept_multiple_files=True)
            
            col1, col2, col3 = st.columns(3)
            with col3:
                if st.button("Gerar Sentença", use_container_width=True, disabled=st.session_state.generating_sentence):
                    st.session_state.generating_sentence = True
                    st.rerun()

            if st.session_state.generating_sentence:
                status_placeholder = col1.empty()
                with status_placeholder:
                    with st.spinner("Gerando sentença..."):
                        try:
                            response = requests.post(
                                f"{API_URL}/gerar-sentenca",
                                data={
                                    "relatorio": st.session_state.relatorio,
                                    "instrucoes_usuario": instrucoes,
                                    "numero_processo": st.session_state.get("numero_processo", ""),
                                    "buscar_na_base": "false"
                                },
                                files=[("arquivos_referencia", (f.name, f.getvalue())) for f in arquivos_ref or []],
                                timeout=600
                            )
                            if response.status_code == 200:
                                result = response.json()
                                st.session_state.sentenca_texto = limpar_relatorio(result.get('sentenca', ''))
                                sent_url = API_URL + result["sentenca_url"]
                                refs_url = API_URL + result["referencias_url"]
                                st.session_state.sentenca_bytes = requests.get(sent_url).content
                                st.session_state.referencias_bytes = requests.get(refs_url).content
                                st.session_state.sentenca_processada = True
                            else:
                                st.error(f"❌ Erro {response.status_code}: {response.text}")
                        except Exception as e:
                            st.error(f"❌ Erro inesperado na geração da sentença: {e}")
                        finally:
                            st.session_state.generating_sentence = False
                            st.rerun()

        if st.session_state.sentenca_processada:
            st.success("⚖️ Sentença gerada! Descarregue os ficheiros.")
            with st.expander("📄 Visualizar Sentença Gerada", expanded=True):
                st.text_area("Conteúdo:", value=st.session_state.sentenca_texto, height=300, disabled=True)

            dl_col1, dl_col2, dl_col3 = st.columns(3)
            nome_sentenca = gerar_nome_arquivo_sentenca(st.session_state.numero_processo)
            
            # CORREÇÃO: Adiciona o parâmetro 'disabled' aos botões
            dl_col3.download_button("Baixar Sentença", data=st.session_state.sentenca_bytes, file_name=nome_sentenca, use_container_width=True, disabled=st.session_state.confirming_clear)
            dl_col2.download_button("Baixar Referências", data=st.session_state.referencias_bytes, file_name="referencias.zip", use_container_width=True, disabled=st.session_state.confirming_clear)

            with dl_col1:
                if st.button("Limpar Dados", use_container_width=True, type="primary", disabled=st.session_state.confirming_clear):
                    st.session_state.confirming_clear = True
                    st.rerun()

            if st.session_state.confirming_clear:
                st.warning("**Tem a certeza de que pretende limpar todos os dados?**")
                col1_clear, col2_clear, _ = st.columns(3)
                if col1_clear.button("Sim, apagar tudo", use_container_width=True, type="primary"):
                    clear_working_data()
                if col2_clear.button("Cancelar", use_container_width=True):
                    st.session_state.confirming_clear = False
                    st.rerun()

        st.markdown("---")
        st.markdown("""
        <div style="text-align: center; color: #888; font-size: 0.9em; margin-top: 2rem;">
            <p><strong>Justino TJPE - Versão BETA 2.2</strong></p>
            <p style="font-weight: bold; color: #ff4b4b;">
                ⚠️ Lembrete: Sempre confira a minuta gerada antes de a lançar no sistema.
            </p>
        </div>
        """, unsafe_allow_html=True)

def run():
    init_session_state()
    if st.session_state.authenticated:
        main_app()
    else:
        show_login_page()

if __name__ == "__main__":
    run()
